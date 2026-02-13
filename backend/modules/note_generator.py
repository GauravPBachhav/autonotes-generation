"""
Note Generator Module
Formats and exports notes to multiple formats (Markdown, PDF, Docx)
Produces two-section output: Original Transcription + Structured Notes
"""

import os
import logging
from typing import Dict, List
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
import markdown

logger = logging.getLogger(__name__)


class NoteGenerator:
    """Generate formatted notes in multiple formats"""

    def __init__(self, output_dir: str = "./output"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    def generate_note_content(
        self,
        title: str,
        transcript_data: Dict,
        processed_data: Dict,
        summaries: Dict,
    ) -> Dict[str, str]:
        """
        Generate structured note content in multiple formats.
        Two main sections: Original Transcription + Structured Notes.
        """
        logger.info(f"Generating note content for: {title}")

        markdown_content = self._build_markdown(
            title, transcript_data, processed_data, summaries,
        )
        html_content = markdown.markdown(markdown_content, extensions=['tables', 'fenced_code'])
        text_content = self._extract_plain_text(markdown_content)

        content = {
            "markdown": markdown_content,
            "html": html_content,
            "text": text_content,
        }

        logger.info("Note content generated successfully")
        return content

    def _build_markdown(
        self,
        title: str,
        transcript_data: Dict,
        processed_data: Dict,
        summaries: Dict,
    ) -> str:
        """Build markdown formatted note with two clear sections"""
        md = []

        # ─── HEADER ───
        md.append(f"# 📚 {title}\n")
        md.append(f"*Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*\n")
        md.append(f"**Duration:** {transcript_data.get('duration', 'N/A')} seconds  ")
        md.append(f"**Language:** {transcript_data.get('language', 'N/A')}  ")
        md.append(f"**Words:** {processed_data.get('word_count', 0)}  ")
        md.append(f"**Topics:** {processed_data.get('section_count', 0)}\n\n")

        # ─── TABLE OF CONTENTS ───
        md.append("## 📋 Table of Contents\n")
        md.append("1. [🎙️ Original Transcription](#-original-transcription)\n")
        md.append("2. [📝 Structured Notes](#-structured-notes)\n")
        md.append("   - [Summary](#summary)\n")
        md.append("   - [Topic-wise Notes](#topic-wise-notes)\n")
        md.append("   - [Key Definitions](#key-definitions)\n")
        md.append("   - [Key Takeaways](#key-takeaways)\n")
        md.append("   - [Quick Revision](#quick-revision)\n")
        md.append("   - [Keywords & Key Phrases](#keywords--key-phrases)\n\n")

        # ═══════════════════════════════════════════════════════════════
        # SECTION 1: ORIGINAL TRANSCRIPTION (as-is)
        # ═══════════════════════════════════════════════════════════════
        md.append("---\n\n")
        md.append("# 🎙️ Original Transcription\n")
        md.append("*Complete word-by-word transcript from audio — preserved as-is*\n\n")

        raw_text = transcript_data.get("text", processed_data.get("original_text", "N/A"))
        md.append(f"{raw_text}\n\n")

        # ═══════════════════════════════════════════════════════════════
        # SECTION 2: STRUCTURED NOTES
        # ═══════════════════════════════════════════════════════════════
        md.append("---\n\n")
        md.append("# 📝 Structured Notes\n")
        md.append("*Well-organized, topic-wise notes generated from the transcription*\n\n")

        # ── Summary ──
        md.append("## Summary\n\n")
        md.append(f"{summaries.get('overall_summary', 'No summary available.')}\n\n")

        # ── Topic-wise Notes ──
        structured = processed_data.get("structured_notes", {})
        topics = structured.get("topics", [])

        md.append("## Topic-wise Notes\n\n")
        if topics:
            for i, topic in enumerate(topics, 1):
                md.append(f"### 📌 {i}. {topic.get('title', f'Topic {i}')}\n\n")

                # Bullet points
                bullets = topic.get("bullet_points", [])
                if bullets:
                    for bullet in bullets:
                        md.append(f"- {bullet}\n")
                else:
                    content = topic.get("content", "")
                    if content:
                        md.append(f"{content}\n")
                md.append("\n")

                # Section keywords
                kws = topic.get("keywords", [])
                if kws:
                    md.append(f"**Key terms:** {', '.join(kws)}\n\n")
        else:
            # Fallback: use sections
            for i, section in enumerate(processed_data.get("sections", []), 1):
                md.append(f"### 📌 {i}. {section.get('title', f'Topic {i}')}\n\n")
                md.append(f"{section.get('text', '')}\n\n")

        # ── Key Definitions ──
        definitions = structured.get("definitions", [])
        if definitions:
            md.append("## Key Definitions\n\n")
            md.append("| Term | Definition |\n")
            md.append("|------|------------|\n")
            for d in definitions:
                term = d.get("term", "")
                defn = d.get("definition", "").replace("|", "\\|")
                md.append(f"| **{term}** | {defn} |\n")
            md.append("\n")

        # ── Key Takeaways ──
        takeaways = structured.get("key_takeaways", [])
        if takeaways:
            md.append("## ⭐ Key Takeaways\n\n")
            for i, point in enumerate(takeaways, 1):
                md.append(f"{i}. {point}\n")
            md.append("\n")

        # ── Bullet Points from Summarizer ──
        bullet_points = summaries.get("bullet_points", [])
        if bullet_points:
            md.append("## 🔑 Important Points\n\n")
            for point in bullet_points:
                md.append(f"- {point}\n")
            md.append("\n")

        # ── Quick Revision ──
        revision = structured.get("quick_revision", [])
        if revision:
            md.append("## 🔄 Quick Revision\n\n")
            for item in revision:
                md.append(f"- {item}\n")
            md.append("\n")

        # ── Keywords & Key Phrases ──
        md.append("## 🏷️ Keywords & Key Phrases\n\n")
        keywords = processed_data.get("keywords", [])
        key_phrases = processed_data.get("key_phrases", [])

        if keywords:
            md.append("**Keywords:** ")
            md.append(", ".join(f"`{kw}`" for kw in keywords[:15]))
            md.append("\n\n")

        if key_phrases:
            md.append("**Key Phrases:** ")
            md.append(", ".join(f"*{kp}*" for kp in key_phrases[:10]))
            md.append("\n\n")

        # ── Statistics ──
        md.append("---\n\n")
        md.append("## 📊 Statistics\n\n")
        md.append(f"| Metric | Value |\n")
        md.append(f"|--------|-------|\n")
        md.append(f"| Word Count | {processed_data.get('word_count', 0)} |\n")
        md.append(f"| Sentence Count | {processed_data.get('sentence_count', 0)} |\n")
        md.append(f"| Topics Detected | {processed_data.get('section_count', 0)} |\n")
        md.append(f"| Duration | {transcript_data.get('duration', 'N/A')}s |\n")
        md.append(f"| Language | {transcript_data.get('language', 'N/A')} |\n")

        return "".join(md)

    def export_markdown(self, content: str, filename: str) -> str:
        """Export note to Markdown file"""
        try:
            output_path = os.path.join(self.output_dir, f"{filename}.md")
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(content)
            
            logger.info(f"Markdown file saved: {output_path}")
            return output_path
        except Exception as e:
            logger.error(f"Failed to export Markdown: {e}")
            raise

    def export_pdf(self, markdown_content: str, filename: str, title: str = "Notes") -> str:
        """Export note to a clean, well-formatted PDF file"""
        try:
            output_path = os.path.join(self.output_dir, f"{filename}.pdf")
            os.makedirs(os.path.dirname(output_path), exist_ok=True)

            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()

            # ── Title ──
            pdf.set_font("Helvetica", "B", 18)
            safe_title = self._strip_emojis(title)
            pdf.cell(0, 12, safe_title, ln=True, align="C")
            pdf.ln(4)

            # ── Process markdown line by line ──
            lines = markdown_content.split("\n")
            page_w = pdf.w - pdf.l_margin - pdf.r_margin  # usable width

            for raw_line in lines:
                line = self._strip_emojis(raw_line).rstrip()

                # Skip empty lines → small space
                if not line.strip():
                    pdf.ln(3)
                    continue

                # Heading 1
                if line.startswith("# "):
                    pdf.ln(4)
                    pdf.set_font("Helvetica", "B", 16)
                    pdf.multi_cell(page_w, 8, line[2:].strip())
                    pdf.ln(2)

                # Heading 2
                elif line.startswith("## "):
                    pdf.ln(3)
                    pdf.set_font("Helvetica", "B", 13)
                    pdf.multi_cell(page_w, 7, line[3:].strip())
                    pdf.ln(2)

                # Heading 3
                elif line.startswith("### "):
                    pdf.ln(2)
                    pdf.set_font("Helvetica", "B", 11)
                    pdf.multi_cell(page_w, 6, line[4:].strip())
                    pdf.ln(1)

                # Table row → format as plain text
                elif line.startswith("|"):
                    cells = [c.strip().strip("*").strip() for c in line.split("|") if c.strip()]
                    if cells and not all(set(c) <= {'-', ' ', ':'} for c in cells):
                        pdf.set_font("Helvetica", "", 10)
                        row_text = "  |  ".join(cells)
                        pdf.multi_cell(page_w, 5, row_text)

                # Horizontal rule
                elif line.strip() == "---":
                    pdf.ln(2)
                    y = pdf.get_y()
                    pdf.line(pdf.l_margin, y, pdf.l_margin + page_w, y)
                    pdf.ln(3)

                # Bullet points
                elif line.strip().startswith("- "):
                    pdf.set_font("Helvetica", "", 10)
                    text = self._clean_md_inline(line.strip()[2:])
                    pdf.multi_cell(page_w, 5, f"  * {text}")

                # Numbered list
                elif len(line.strip()) > 2 and line.strip()[0].isdigit() and line.strip()[1] in '.):':
                    pdf.set_font("Helvetica", "", 10)
                    text = self._clean_md_inline(line.strip())
                    pdf.multi_cell(page_w, 5, f"  {text}")

                # Italic metadata
                elif line.strip().startswith("*") and line.strip().endswith("*"):
                    pdf.set_font("Helvetica", "I", 9)
                    text = line.strip().strip("*")
                    pdf.multi_cell(page_w, 5, text)

                # Regular paragraph
                else:
                    pdf.set_font("Helvetica", "", 10)
                    text = self._clean_md_inline(line)
                    if text.strip():
                        pdf.multi_cell(page_w, 5, text)

            pdf.output(output_path)

            logger.info(f"PDF file saved: {output_path}")
            return output_path
        except Exception as e:
            logger.error(f"Failed to export PDF: {e}")
            raise

    @staticmethod
    def _strip_emojis(text: str) -> str:
        """Remove emoji and other non-Latin1 characters that break PDF fonts."""
        import re
        # Remove characters outside BMP Latin range that built-in fonts can't render
        text = re.sub(
            r'[\U00010000-\U0010ffff]'   # supplementary planes (most emojis)
            r'|[\u2600-\u27BF]'           # misc symbols
            r'|[\uFE00-\uFE0F]'           # variation selectors
            r'|[\u2B50-\u2B55]'           # stars etc
            r'|[\u23CF-\u23FA]'           # misc technical
            r'|[\u200D]'                  # zero-width joiner
            r'|[\u2702-\u27B0]'           # dingbats
            r'|[\U0001F000-\U0001FFFF]',  # catch-all emoticons/symbols
            '', text
        )
        # Replace common unicode chars that break Latin-1 PDF fonts
        text = text.replace('\u2014', '-')   # em-dash —
        text = text.replace('\u2013', '-')   # en-dash –
        text = text.replace('\u2018', "'")   # left single quote '
        text = text.replace('\u2019', "'")   # right single quote '
        text = text.replace('\u201C', '"')   # left double quote "
        text = text.replace('\u201D', '"')   # right double quote "
        text = text.replace('\u2026', '...')  # ellipsis …
        text = text.replace('\u2022', '*')   # bullet •
        text = text.replace('\u00A0', ' ')   # non-breaking space
        text = text.replace('\u2192', '->')  # arrow →
        text = text.replace('\u2190', '<-')  # arrow ←
        # Strip any remaining non-Latin-1 characters as last resort
        text = text.encode('latin-1', errors='replace').decode('latin-1')
        # Clean up double spaces
        text = re.sub(r'  +', ' ', text)
        return text.strip()

    @staticmethod
    def _clean_md_inline(text: str) -> str:
        """Remove inline markdown formatting (bold, italic, code)."""
        text = text.replace("**", "").replace("__", "")
        text = text.replace("*", "").replace("_", "")
        text = text.replace("`", "")
        return text

    def export_docx(self, markdown_content: str, filename: str, title: str = "Notes") -> str:
        """Export note to DOCX file"""
        try:
            output_path = os.path.join(self.output_dir, f"{filename}.docx")
            os.makedirs(os.path.dirname(output_path), exist_ok=True)

            doc = Document()

            # Title
            safe_title = self._strip_emojis(title)
            title_para = doc.add_heading(safe_title, level=1)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Timestamp
            doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

            # Parse markdown into structured DOCX elements
            lines = markdown_content.split("\n")
            for raw_line in lines:
                line = self._strip_emojis(raw_line).rstrip()

                if not line.strip():
                    continue

                if line.startswith("# "):
                    doc.add_heading(line[2:].strip(), level=1)
                elif line.startswith("## "):
                    doc.add_heading(line[3:].strip(), level=2)
                elif line.startswith("### "):
                    doc.add_heading(line[4:].strip(), level=3)
                elif line.startswith("- "):
                    text = self._clean_md_inline(line.strip()[2:])
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(text)
                elif line.strip().startswith("|") and not all(c in '-| :' for c in line.strip()):
                    cells = [c.strip().strip("*") for c in line.split("|") if c.strip()]
                    doc.add_paragraph("  |  ".join(cells))
                elif line.strip() == "---":
                    continue
                else:
                    text = self._clean_md_inline(line)
                    if text.strip():
                        doc.add_paragraph(text)

            doc.save(output_path)

            logger.info(f"DOCX file saved: {output_path}")
            return output_path
        except Exception as e:
            logger.error(f"Failed to export DOCX: {e}")
            raise

    def _extract_plain_text(self, markdown_content: str) -> str:
        """Extract plain text from markdown"""
        # Remove markdown syntax
        text = markdown_content
        text = text.replace("# ", "").replace("## ", "").replace("### ", "")
        text = text.replace("**", "").replace("*", "").replace("_", "")
        text = text.replace("[", "").replace("]", "").replace("(", "").replace(")", "")
        return text

    def export_all_formats(
        self,
        note_content: Dict,
        filename: str,
        title: str = "Lecture Notes",
    ) -> Dict[str, str]:
        """
        Export note to all available formats
        
        Returns:
            Dictionary with paths to all exported files
        """
        logger.info(f"Exporting all formats for: {filename}")
        
        markdown_path = self.export_markdown(note_content["markdown"], filename)
        pdf_path = self.export_pdf(note_content["markdown"], filename, title)
        docx_path = self.export_docx(note_content["markdown"], filename, title)

        result = {
            "markdown": markdown_path,
            "pdf": pdf_path,
            "docx": docx_path,
        }

        logger.info(f"All formats exported successfully")
        return result
